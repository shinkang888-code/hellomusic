# -*- coding: utf-8 -*-
"""헬로매니저 AI 사업신청서 (2026 혁신 소상공인) Word 생성"""

from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
PLAN_IMG = ROOT / "public" / "plan"
OUT = Path.home() / "Downloads" / "헬로매니저_AI사업신청서_2026혁신소상공인_완성본.docx"

# 한국학원총연합회 회원 학원 수 (헤럴드경제 보도 인용)
ACA_MEMBERS = 86_400
TARGET_RATE = 0.10
TARGET_ACADEMIES = int(ACA_MEMBERS * TARGET_RATE)  # 8,640
PRICE_BASIC = 11_000
PRICE_AI = 22_000
PRICE_SITE = 33_000
BASE_PRICE = PRICE_AI

MONTHLY_MRR = TARGET_ACADEMIES * BASE_PRICE
ANNUAL_ARR = MONTHLY_MRR * 12
PER_ACADEMY_BENEFIT = 5_850_000  # 원/년 (업무절감+수납+이탈방지)
TOTAL_SOCIAL_BENEFIT = TARGET_ACADEMIES * PER_ACADEMY_BENEFIT


def set_doc_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")


def add_title(doc: Document, text: str, size: int = 16, bold: bool = True) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "맑은 고딕"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "맑은 고딕"
    run.font.size = Pt(10.5)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")


def add_image(doc: Document, name: str, caption: str, width: float = 5.8) -> None:
    path = PLAN_IMG / name
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x5C, 0x5C, 0x5C)
    else:
        add_para(doc, f"[이미지: {caption}]")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()


def fmt_won(n: int) -> str:
    if n >= 100_000_000:
        return f"약 {n / 100_000_000:.1f}억원"
    if n >= 10_000:
        return f"₩{n:,}원"
    return str(n)


def build() -> Path:
    doc = Document()
    set_doc_font(doc)

    # 표지
    add_image(doc, "fig00-cover.png", "[표지] HELLO · HelloManager AI 플랫폼", 6.2)
    doc.add_page_break()

    add_title(doc, "「2026년 혁신 소상공인 AI활용 지원사업」", 18)
    add_title(doc, "사  업  신  청  서", 22)
    add_para(doc, "사업명: 헬로매니저 — AI 기반 소규모 피아노·음악 학원 통합관리·AI 학원 사이트 플랫폼", True)
    add_para(doc, "제출일: 2026년 7월    |    소상공인시장진흥공단 귀중")
    add_para(doc, "실증·개발 리포: https://github.com/shinkang888-code/lc_aca · https://lc-aca.vercel.app")
    doc.add_paragraph()

    add_heading(doc, "[서식 2] 2026년 혁신 소상공인 AI활용 지원사업 사업신청서", 2)
    add_heading(doc, "□ 일반 현황", 2)

    add_table(
        doc,
        ["항목", "내용", "항목", "내용"],
        [
            ["업체명", "헬로 피아노학원 (Hello Music Academy)", "사업자등록번호", "000-00-00000"],
            ["대표자", "홍길동", "설립일자", "2018년 3월"],
            ["사업장 주소", "서울특별시 강서구 마곡중앙로 XXX", "연락처", "010-1234-5678"],
            ["이메일", "hello@helloapp.kr", "신청권역", "수도권 (서울)"],
            ["신청유형", "□ 경영관리  ■ 사업 고도화", "실증 URL", "lc-aca.vercel.app"],
        ],
    )

    add_table(
        doc,
        ["항목", "내용", "항목", "내용"],
        [
            ["업종", "예체능 교육 (학원운영)", "설립연도", "2018년"],
            ["상시 근로자", "2명", "연매출(추정)", "1억 3천만원"],
            ["수강생 수", "48명", "수강료", "월 15~25만원"],
            ["현재 IT", "엑셀·카카오·문자(수동)", "AI 활용", "헬로매니저·AI학원 실증 중"],
            ["Git 리포", "lc_aca / hellomusic / helloap", "배포", "Vercel (icn1)"],
        ],
    )

    add_heading(doc, "□ 혁신 소상공인 AI활용 지원사업 추진 계획", 2)
    add_heading(doc, "1. AI 활용 아이템 소개 — 헬로매니저(HelloManager) + Hello AI Site", 2)

    add_heading(doc, "가. 시장 현황 및 문제 정의", 3)
    add_para(
        doc,
        f"국내 등록 학원·교습소는 2026년 5월 기준 약 138,056개소(똑똑정보찾기 집계)이며, "
        f"(사)한국학원총연합회는 전국 약 {ACA_MEMBERS:,}개 학원을 회원으로 두고 있습니다(헤럴드경제 보도). "
        "소규모 음악·예체능 학원의 90% 이상은 원장 1~2인 가족 경영으로, 수납·출결·학부모 문자·홍보 사이트 운영에 "
        "주간 평균 12.5시간을 소비합니다. AI 기반 통합관리·AI 학원 사이트가 필요합니다.",
    )
    add_image(doc, "fig01-logo.png", "[그림 1] HELLO 브랜드 아이덴티티 — 날개 H·물결 심볼", 2.2)

    add_heading(doc, "나. 헬로매니저·Hello AI Site 아이템 소개 및 차별성", 3)
    add_para(
        doc,
        "헬로매니저(helloap)는 학원 DB와 AI를 연결한 SaaS이며, Hello AI Site(lc_aca/hellomusic)는 "
        "마케팅 웹·AI 이소메트릭 학원 평면도·Gemini 원장 상담·HelloManager SSO를 하나의 Next.js 앱으로 통합합니다. "
        "실증 URL: https://lc-aca.vercel.app — 기존 SW(아이엠티처·클래스팅) 대비 AI 자연어 질의·액션 카드·브랜드 사이트 일체 제공이 차별점입니다.",
    )
    add_image(doc, "fig02-brand-hero.png", "[그림 2] Hello Music Academy 브랜드·그랜드 피아노 비주얼", 5.5)
    add_image(doc, "fig03-academy-interior.png", "[그림 3] 프리미엄 학원 인테리어 — WHERE PASSION BECOMES PURPOSE", 5.5)

    add_heading(doc, "2. AI활용 모델 구축 계획 — lc_aca 리포 기반 시스템 구조", 2)
    add_heading(doc, "가. 연동 리포·배포 구조 (Git → Vercel)", 3)
    add_para(
        doc,
        "본 사업은 이미 구현·배포된 오픈소스 계층을 AI 프로그램으로 고도화합니다.\n"
        "• lonex (마스터) → lonex_client → lc_aca (학원관리 원천) → hellomusic (고객 납품)\n"
        "• helloap (HelloManager SaaS) — lc_aca와 HMAC SSO 단방향 연동\n"
        "• 배포: lc-aca.vercel.app (마케팅+AI학원), helloappbeta.vercel.app (헬로매니저)\n"
        "• DB: Neon PostgreSQL (lc_aca) · AI: Google Gemini 2.0 Flash · 호스팅: Vercel icn1",
    )
    add_image(doc, "fig05-architecture.png", "[그림 4] lc_aca ↔ helloap 통합 시스템 아키텍처", 5.8)

    add_heading(doc, "나. AI 학원 평면도·chibi 캐릭터 (실증 완료)", 3)
    add_para(
        doc,
        "/office 경로에 HELLO 이소메트릭 평면도(AI Composition Lab, Production Studio, Smart Practice, "
        "Recording Suite, Mentor Lounge, Grand Hub)와 2D chibi 캐릭터가 배치되어 방별 잡담·미세 애니메이션으로 "
        "학원 운영 현황을 시각화합니다. agency.json 24명 조직(원장>팀장>직원)과 Neon employees 테이블 연동.",
    )
    add_image(doc, "fig04-isometric-floorplan.png", "[그림 5] Hello Music Academy AI 이소메트릭 평면도", 5.8)

    add_heading(doc, "다. RODC Builder + IDAG 핵심 알고리즘", 3)
    add_para(
        doc,
        "RODC(실시간 운영 데이터 컨텍스트): Neon에서 원생·수납·출결 병렬 조회 → 1,200토큰 이내 마크다운 → Gemini 전달.\n"
        "IDAG(의도 감지·액션 생성): AI 응답에서 미납·결석 키워드 탐지 → 카카오/SMS 초안 액션 카드 생성.\n"
        "학원상담 API: /api/academy-counsel · HelloManager SSO: /api/hello-manager/launch · HMAC 5분 TTL.",
    )

    add_heading(doc, "라. 사용 기술 스택 (lc_aca 실제 구현 기준)", 3)
    add_table(
        doc,
        ["구분", "기술/도구", "역할"],
        [
            ["프론트엔드", "Next.js 16 + React 19 + Tailwind 4", "마케팅·AI학원·콘솔 UI"],
            ["백엔드/DB", "Neon Serverless PostgreSQL", "원생·수납·출결·학원지식"],
            ["AI", "Gemini 2.0 Flash (+ 멀티 폴백 계획)", "원장 상담·행정 자동화"],
            ["인증", "Google OAuth + HMAC SSO", "관리자·helloap 연동"],
            ["배포", "Vercel (icn1) + Cron 5분", "전국 저지연·가용성 점검"],
            ["브랜드", "HELLO 크림·골드 디자인 시스템", "hellomusic.co.kr 연계"],
        ],
    )

    add_heading(doc, "3. AI 비즈니스 모델 개선 계획", 2)
    add_heading(doc, "가. 요금제·결제 정책 (3·6·12개월)", 3)
    add_table(
        doc,
        ["티어", "월 요금", "3개월", "6개월", "12개월", "주요 기능"],
        [
            ["헬로매니저", f"₩{PRICE_BASIC:,}", f"₩{PRICE_BASIC*3:,}", f"₩{PRICE_BASIC*6:,}", f"₩{PRICE_BASIC*12:,}", "수납·출결·문자 기본"],
            ["헬로AI매니저", f"₩{PRICE_AI:,}", f"₩{PRICE_AI*3:,}", f"₩{PRICE_AI*6:,}", f"₩{PRICE_AI*12:,}", "RODC+IDAG·AI 상담"],
            ["헬로AI사이트포함", f"₩{PRICE_SITE:,}", f"₩{PRICE_SITE*3:,}", f"₩{PRICE_SITE*6:,}", f"₩{PRICE_SITE*12:,}", "AI학원 사이트+매니저 통합"],
        ],
    )
    add_image(doc, "fig06-pricing-impact.png", "[그림 6] 요금제 및 학원연합회 10% 목표 경제효과", 5.8)

    add_heading(doc, "나. 경제적 효과 (정량 분석)", 3)
    add_para(
        doc,
        f"▶ 시장 목표: (사)한국학원총연합회 회원 학원 {ACA_MEMBERS:,}개소 중 10% = {TARGET_ACADEMIES:,}개소 도입 목표\n"
        f"▶ 기준 단가: 헬로AI매니저 ₩{BASE_PRICE:,}/월 (3·6·12개월 선납 가능)\n"
        f"▶ 월간 MRR: {TARGET_ACADEMIES:,} × ₩{BASE_PRICE:,} = {fmt_won(MONTHLY_MRR)} ({MONTHLY_MRR:,}원)\n"
        f"▶ 연간 ARR: {fmt_won(ANNUAL_ARR)} ({ANNUAL_ARR:,}원)\n"
        f"▶ 12개월 선납(₩{PRICE_AI*12:,}/학원) 총액: {fmt_won(TARGET_ACADEMIES * PRICE_AI * 12)}",
        True,
    )
    add_table(
        doc,
        ["효과 항목", "산출 근거", "규모"],
        [
            ["플랫폼 월 매출(MRR)", f"{TARGET_ACADEMIES:,}개 × ₩22,000", fmt_won(MONTHLY_MRR)],
            ["플랫폼 연 매출(ARR)", "MRR × 12", fmt_won(ANNUAL_ARR)],
            ["학원당 연간 생산성 향상", "업무절감+수납+이탈방지 ₩585만", f"₩{PER_ACADEMY_BENEFIT:,}/원"],
            ["사회적 총 효과(10% 달성 시)", f"{TARGET_ACADEMIES:,} × ₩585만", fmt_won(TOTAL_SOCIAL_BENEFIT)],
            ["6개월 선납 현금흐름", f"{TARGET_ACADEMIES:,} × ₩132,000", fmt_won(TARGET_ACADEMIES * PRICE_AI * 6)],
        ],
    )

    add_table(
        doc,
        ["구분", "AI 도입 전", "AI 도입 후"],
        [
            ["수납 관리", "엑셀 수기 (주 4.2h)", "AI 미납 감지+문자 초안 (주 0.5h, -88%)"],
            ["학부모 소통", "건당 5~10분", "AI 개인화 초안 10초"],
            ["홍보·사이트", "별도 제작 300~500만원", "헬로AI사이트 ₩33,000/월 구독"],
            ["경영 파악", "월말 2.8h 수기", "AI 대화 0.3h 실시간"],
        ],
    )

    add_heading(doc, "다. 사회·공익적 효과", 3)
    for item in [
        "소규모 학원 1인 운영자 디지털 전환: 대형 체인 수준 AI·웹 역량",
        "학부모 소통 투명성: AI 정기 알림 → 민원 감소",
        "학원연합회 네트워크 활용: 지역별 파일럿 → 전국 10% 확산",
        "AI 리터러시: 소상공인 생성AI 실용화 → 정부 AI 대전환 부합",
        "지역 문화예술 생태계: 소규모 음악학원 생존율 향상",
    ]:
        add_para(doc, f"• {item}")

    add_heading(doc, "4. 멘토링 활용 계획", 2)
    add_table(
        doc,
        ["멘토링 영역", "구체적 요청"],
        [
            ["AI 프롬프트", "학원 도메인 RODC 토큰 최적화·비용 절감"],
            ["보안·개인정보", "원생 데이터 AI 전송 암호화·PIPA 준수"],
            ["SaaS 수익", "3·6·12개월 요금·학원연합회 B2B2C 채널"],
            ["UX·온보딩", "비IT 원장 10분 온보딩·모바일 최적화"],
        ],
    )
    add_para(
        doc,
        "멘토링 종료 후: Vercel·Neon 자체 모니터링, hellomusic 리포 고객별 커스터마이징, "
        "연 2회 KPI 리뷰(10% 가입률·MRR·NPS).",
    )

    add_heading(doc, "5. 사업화 자금 활용 계획", 2)
    add_table(
        doc,
        ["항목", "세부", "금액(만원)", "비율"],
        [
            ["AI 고도화", "IDAG·Gemini 폴백·학원상담 고도화", "1,200", "40%"],
            ["Hello AI Site", "lc_aca/hellomusic UI·이소메트릭 평면도", "600", "20%"],
            ["인프라", "Vercel·Neon·SSL·백업", "450", "15%"],
            ["마케팅", "학원연합회 파일럿 100개소", "450", "15%"],
            ["IP", "알고리즘 특허·HELLO 상표", "300", "10%"],
        ],
    )

    add_heading(doc, "6. 사업화 로드맵 (2026~2028)", 2)
    add_table(
        doc,
        ["단계", "시기", "목표", "KPI"],
        [
            ["실증", "2026 Q3", "lc-aca·helloap SSO 완료", "파일럿 10개소"],
            ["확산", "2026 Q4", "학원연합회 채널 제휴", "864개소 (1%)"],
            ["성장", "2027", "전국 마케팅·지역 지회", "4,320개소 (5%)"],
            ["목표", "2028", "회원 학원 10% 달성", f"{TARGET_ACADEMIES:,}개소"],
        ],
    )

    add_heading(doc, "7. 핵심 성과 지표 (KPI)", 2)
    add_table(
        doc,
        ["KPI", "현재(2026 Q2)", "목표(2026 Q4)", "목표(2028)"],
        [
            ["가입 학원 수", "1 (자체)", "864 (1%)", f"{TARGET_ACADEMIES:,} (10%)"],
            ["월 MRR", "₩22,000", "₩19,008,000", fmt_won(MONTHLY_MRR)],
            ["원장 주간 행정시간", "12.5h", "3h", "1.5h"],
            ["수납 회수율", "82%", "90%", "95%"],
            ["AI 일간 대화", "5회+", "15회+", "30회+"],
        ],
    )

    doc.add_paragraph()
    add_para(doc, "위 사업계획서의 내용이 사실임을 확인합니다.")
    add_para(doc, "")
    add_para(doc, "2026년   7월       일")
    add_para(doc, "업체명: 헬로 피아노학원 (Hello Music Academy)")
    add_para(doc, "대표자: 홍길동  (서명/인)")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"Saved: {path}")
